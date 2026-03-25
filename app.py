"""
AI Resume Analyser — Single File Edition
Powered by Groq API (LLaMA 3.3 70B) + Plotly visualisations

Run:  streamlit run app.py
"""

import io
import os
import json
import re
from datetime import datetime

import pdfplumber
import docx
import plotly.graph_objects as go
import streamlit as st
from dotenv import load_dotenv
from groq import Groq

load_dotenv()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 1 — JOB ROLES
# ═══════════════════════════════════════════════════════════════════════════════

JOB_ROLES = {
    "Penetration Tester / Ethical Hacker": {
        "icon": "🔴",
        "description": "Tests systems, networks, and applications for vulnerabilities using attacker methodologies.",
        "required_skills": [
            "Python", "Linux/Kali Linux", "Metasploit", "Burp Suite", "Nmap",
            "OWASP Top 10", "Network Security", "Web Application Security", "SQL Injection"
        ],
        "preferred_skills": [
            "OSCP", "CEH", "Bug Bounty", "CTF Experience", "Exploit Development",
            "Active Directory", "Social Engineering", "Wireless Security", "Reverse Engineering"
        ],
        "soft_skills_needed": ["Report Writing", "Critical Thinking", "Problem Solving"],
        "weights": {
            "technical_skills": 0.28,
            "cybersecurity_relevance": 0.35,
            "experience": 0.20,
            "education": 0.07,
            "soft_skills": 0.05,
            "ats_compatibility": 0.05
        },
        "hire_threshold": 62,
        "min_experience_years": 1
    },
    "Security Analyst (SOC / DFIR)": {
        "icon": "🔵",
        "description": "Monitors, detects, and responds to security incidents from a Security Operations Center.",
        "required_skills": [
            "SIEM", "Splunk", "Incident Response", "Threat Intelligence", "Log Analysis",
            "Vulnerability Assessment", "Malware Analysis", "Wireshark", "TCP/IP"
        ],
        "preferred_skills": [
            "CompTIA Security+", "CEH", "CISSP", "CySA+", "ELK Stack",
            "Digital Forensics", "MITRE ATT&CK", "Threat Hunting", "Python"
        ],
        "soft_skills_needed": ["Analytical Thinking", "Communication", "Documentation"],
        "weights": {
            "technical_skills": 0.25,
            "cybersecurity_relevance": 0.32,
            "experience": 0.22,
            "education": 0.10,
            "soft_skills": 0.06,
            "ats_compatibility": 0.05
        },
        "hire_threshold": 60,
        "min_experience_years": 1
    },
    "DevSecOps Engineer": {
        "icon": "🟢",
        "description": "Integrates security into the CI/CD pipeline and software development lifecycle.",
        "required_skills": [
            "Docker", "Kubernetes", "CI/CD", "Jenkins/GitHub Actions", "Python or Bash",
            "SAST/DAST Tools", "AWS / Azure / GCP", "Git", "Infrastructure as Code"
        ],
        "preferred_skills": [
            "Terraform", "Ansible", "OWASP", "Snyk", "SonarQube",
            "HashiCorp Vault", "Linux", "Kubernetes Security", "Zero Trust"
        ],
        "soft_skills_needed": ["Collaboration", "Communication", "Agile Mindset"],
        "weights": {
            "technical_skills": 0.35,
            "cybersecurity_relevance": 0.25,
            "experience": 0.22,
            "education": 0.08,
            "soft_skills": 0.05,
            "ats_compatibility": 0.05
        },
        "hire_threshold": 65,
        "min_experience_years": 2
    },
    "Full Stack Developer (Security-Focused)": {
        "icon": "🟡",
        "description": "Builds secure web applications with frontend and backend expertise and security awareness.",
        "required_skills": [
            "JavaScript", "React / Vue / Angular", "Node.js / Python / Java", "REST APIs",
            "SQL / NoSQL Databases", "Git", "HTML/CSS", "Secure Coding Practices"
        ],
        "preferred_skills": [
            "OWASP Top 10", "OAuth / JWT", "Input Validation", "HTTPS/TLS",
            "Docker", "TypeScript", "Unit/Integration Testing", "Agile/Scrum"
        ],
        "soft_skills_needed": ["Problem Solving", "Team Collaboration", "Communication"],
        "weights": {
            "technical_skills": 0.40,
            "cybersecurity_relevance": 0.18,
            "experience": 0.22,
            "education": 0.10,
            "soft_skills": 0.05,
            "ats_compatibility": 0.05
        },
        "hire_threshold": 63,
        "min_experience_years": 1
    },
    "AI / ML Engineer": {
        "icon": "🤖",
        "description": "Develops and deploys machine learning models with awareness of adversarial attacks and AI safety.",
        "required_skills": [
            "Python", "TensorFlow or PyTorch", "Machine Learning", "Deep Learning",
            "Data Preprocessing", "Model Evaluation", "NumPy / Pandas", "Scikit-learn"
        ],
        "preferred_skills": [
            "LLMs / Transformers", "MLOps", "Docker", "SQL", "Computer Vision",
            "NLP", "Adversarial ML", "Federated Learning", "Model Explainability"
        ],
        "soft_skills_needed": ["Research Aptitude", "Mathematical Thinking", "Documentation"],
        "weights": {
            "technical_skills": 0.42,
            "cybersecurity_relevance": 0.12,
            "experience": 0.22,
            "education": 0.14,
            "soft_skills": 0.05,
            "ats_compatibility": 0.05
        },
        "hire_threshold": 65,
        "min_experience_years": 0
    },
    "Network Security Engineer": {
        "icon": "🌐",
        "description": "Designs and maintains secure network infrastructure for enterprise environments.",
        "required_skills": [
            "Cisco / Juniper", "Firewalls", "VPN", "IDS / IPS", "Network Protocols",
            "TCP/IP", "Routing & Switching", "VLAN", "Network Monitoring"
        ],
        "preferred_skills": [
            "CCNA / CCNP", "Palo Alto Networks", "Fortinet", "Zero Trust Architecture",
            "SD-WAN", "BGP / OSPF", "Python for Automation", "NSE Certifications"
        ],
        "soft_skills_needed": ["Problem Solving", "Documentation", "Vendor Management"],
        "weights": {
            "technical_skills": 0.30,
            "cybersecurity_relevance": 0.30,
            "experience": 0.22,
            "education": 0.08,
            "soft_skills": 0.05,
            "ats_compatibility": 0.05
        },
        "hire_threshold": 63,
        "min_experience_years": 2
    },
    "Cloud Security Engineer": {
        "icon": "☁️",
        "description": "Secures cloud infrastructure and enforces compliance across AWS, Azure, or GCP environments.",
        "required_skills": [
            "AWS / Azure / GCP", "IAM", "Cloud Security Posture Management", "Encryption",
            "Security Groups / NACLs", "VPC Networking", "Compliance (SOC2/PCI/ISO 27001)", "Python or Terraform"
        ],
        "preferred_skills": [
            "AWS Security Specialty", "Azure Security Engineer", "CCSP", "CSPM Tools",
            "Prisma Cloud / Wiz", "Kubernetes Security", "Zero Trust", "SAST/DAST"
        ],
        "soft_skills_needed": ["Risk Assessment", "Communication", "Documentation"],
        "weights": {
            "technical_skills": 0.30,
            "cybersecurity_relevance": 0.28,
            "experience": 0.24,
            "education": 0.08,
            "soft_skills": 0.05,
            "ats_compatibility": 0.05
        },
        "hire_threshold": 65,
        "min_experience_years": 2
    }
}


def get_role_info(role_name: str) -> dict:
    return JOB_ROLES.get(role_name, {})


def get_role_list() -> list:
    return list(JOB_ROLES.keys())


# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 2 — RESUME PARSER
# ═══════════════════════════════════════════════════════════════════════════════

def extract_text_from_resume(uploaded_file) -> str:
    """
    Extract raw text from an uploaded resume file.
    Supports PDF, DOCX, and TXT. Returns cleaned text string.
    """
    filename = uploaded_file.name.lower()
    file_bytes = uploaded_file.read()

    if not file_bytes:
        raise ValueError("The uploaded file is empty. Please upload a valid resume.")

    if filename.endswith(".pdf"):
        return _extract_from_pdf(file_bytes)
    elif filename.endswith(".docx"):
        return _extract_from_docx(file_bytes)
    elif filename.endswith(".txt"):
        return _extract_from_txt(file_bytes)
    else:
        ext = filename.rsplit(".", 1)[-1].upper() if "." in filename else "UNKNOWN"
        raise ValueError(f"Unsupported file format: .{ext}\nPlease upload a PDF, DOCX, or TXT file.")


def _extract_from_pdf(file_bytes: bytes) -> str:
    text_parts = []
    try:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            if not pdf.pages:
                raise ValueError("This PDF has no readable pages.")
            for page in pdf.pages:
                text = page.extract_text()
                if text and text.strip():
                    text_parts.append(text.strip())
    except ValueError:
        raise
    except Exception as e:
        raise RuntimeError(f"Failed to read PDF: {str(e)}\nEnsure the file is not password-protected or corrupted.")

    if not text_parts:
        raise ValueError("No text could be extracted from this PDF.\nIt may be a scanned/image PDF. Please use a text-based PDF.")

    return "\n\n".join(text_parts)


def _extract_from_docx(file_bytes: bytes) -> str:
    paragraphs = []
    try:
        doc = docx.Document(io.BytesIO(file_bytes))
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)
        for table in doc.tables:
            for row in table.rows:
                row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_cells:
                    paragraphs.append("  |  ".join(row_cells))
    except Exception as e:
        raise RuntimeError(f"Failed to read DOCX: {str(e)}\nEnsure the file is a valid .docx format.")

    if not paragraphs:
        raise ValueError("No text could be extracted from this DOCX file.")

    return "\n".join(paragraphs)


def _extract_from_txt(file_bytes: bytes) -> str:
    for encoding in ("utf-8", "utf-16", "latin-1", "cp1252"):
        try:
            text = file_bytes.decode(encoding).strip()
            if text:
                return text
        except (UnicodeDecodeError, LookupError):
            continue
    raise RuntimeError("Could not decode the text file.\nPlease save it as UTF-8 and try again.")


def get_word_count(text: str) -> int:
    return len(text.split())


# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 3 — AI ANALYZER
# ═══════════════════════════════════════════════════════════════════════════════

GROQ_MODEL = "llama-3.3-70b-versatile"

SYSTEM_PROMPT = """You are Dr. Alexandra Reid, a senior HR analyst and cybersecurity talent acquisition specialist \
with 20 years of experience hiring at top-tier firms including CrowdStrike, Palo Alto Networks, and IBM X-Force.

Your evaluations are:
- Evidence-based: you reference specific content from the resume
- Brutally honest: no generic praise, no softening
- Specific: every weakness includes a concrete improvement path
- Calibrated: you score consistently against industry benchmarks

CRITICAL RULE: You must respond ONLY with a valid JSON object.
Do NOT include any preamble, explanation, or markdown code fences.
Your response must start with { and end with } — nothing else."""


def build_prompt(resume_text: str, job_role: str, job_description: str, role_info: dict) -> str:
    required = ", ".join(role_info.get("required_skills", []))
    preferred = ", ".join(role_info.get("preferred_skills", []))
    jd_block = (
        f"\nEMPLOYER JOB DESCRIPTION:\n{job_description.strip()}"
        if job_description and job_description.strip()
        else f"\nNo JD provided — evaluate against standard industry benchmarks for {job_role}."
    )
    resume_excerpt = resume_text[:7000] if len(resume_text) > 7000 else resume_text

    return f"""You are evaluating a candidate for: {job_role}

REQUIRED SKILLS: {required}
PREFERRED/BONUS SKILLS: {preferred}
{jd_block}

RESUME:
===
{resume_excerpt}
===

Return a single JSON object with EXACTLY this schema (all fields required, no extras):

{{
  "candidate_name": "Full name or 'Not Specified'",
  "email": "email address or null",
  "phone": "phone number or null",
  "location": "city/country or null",
  "linkedin": "LinkedIn URL or null",
  "github": "GitHub URL or null",
  "portfolio": "portfolio URL or null",

  "professional_summary": "Your 2-3 sentence honest assessment specific to {job_role}. Reference the actual resume content.",

  "scores": {{
    "technical_skills": <integer 0-100>,
    "cybersecurity_relevance": <integer 0-100>,
    "experience": <integer 0-100>,
    "education": <integer 0-100>,
    "soft_skills": <integer 0-100>,
    "ats_compatibility": <integer 0-100>
  }},

  "overall_score": <integer 0-100, not a simple average — weight by role importance>,
  "experience_years": <number, use 0 for students with no full-time experience>,
  "education_level": "High School | Diploma/Associate | Bachelor's | Master's | PhD | Not Specified",
  "education_field": "Computer Science | Information Technology | Cybersecurity | ECE | Other | Not Specified",
  "institution_name": "University or college name, or Not Specified",
  "projects_count": <integer count of distinct projects mentioned>,

  "technical_skills_found": ["skill1", "skill2"],
  "cybersecurity_skills_found": ["sec-skill1", "sec-skill2"],
  "soft_skills_found": ["soft-skill1"],
  "certifications_found": ["cert1"],
  "tools_found": ["tool1", "tool2"],
  "programming_languages": ["lang1", "lang2"],

  "required_skills_missing": ["skills required for {job_role} that are NOT mentioned in the resume"],
  "recommended_certifications": ["top 3 certs this person should pursue for {job_role}"],

  "strengths": [
    {{
      "title": "3-5 word title",
      "description": "Specific, evidence-based explanation. Quote or paraphrase actual resume content."
    }}
  ],

  "weaknesses": [
    {{
      "title": "3-5 word title",
      "description": "Specific gap and its impact on the candidacy for {job_role}. Be direct."
    }}
  ],

  "red_flags": ["Specific concern: gaps, vague claims, mismatched dates. Empty list if none."],
  "green_flags": ["Specific impressive element: certifications, standout projects, publications."],

  "improvements": [
    {{
      "priority": "HIGH | MEDIUM | LOW",
      "area": "e.g. Missing Certifications | Project Descriptions | Keyword Optimization",
      "current_state": "What the resume shows now (be specific)",
      "suggestion": "Specific, actionable step they should take"
    }}
  ],

  "interview_questions": [
    "Question to probe a specific weakness or verify a claim from this resume",
    "Technical depth question specific to {job_role}",
    "Behavioral or situational question based on their background"
  ],

  "hire_recommendation": "HIRE | MAYBE | REJECT",
  "hire_confidence": <integer 0-100>,
  "hire_reasoning": "3-4 sentences. Be specific — name actual resume elements that drove your decision."
}}"""


def analyze_resume(resume_text: str, job_role: str, job_description: str, api_key: str, role_info: dict) -> dict:
    try:
        client = Groq(api_key=api_key)
    except Exception as e:
        raise RuntimeError(f"Could not initialise Groq client: {str(e)}")

    prompt = build_prompt(resume_text, job_role, job_description, role_info)

    try:
        response = client.chat.completions.create(
            model=GROQ_MODEL,
            messages=[
                {"role": "system", "content": SYSTEM_PROMPT},
                {"role": "user", "content": prompt},
            ],
            temperature=0.2,
            max_tokens=4096,
        )
    except Exception as e:
        raise RuntimeError(f"Groq API call failed: {str(e)}\nCheck your API key and internet connection.")

    return _parse_and_validate(response.choices[0].message.content)


def _parse_and_validate(raw_text: str) -> dict:
    clean = raw_text.strip()
    clean = re.sub(r"^```(?:json)?\s*", "", clean, flags=re.IGNORECASE)
    clean = re.sub(r"\s*```\s*$", "", clean).strip()

    try:
        data = json.loads(clean)
    except json.JSONDecodeError:
        match = re.search(r"\{.*\}", clean, re.DOTALL)
        if match:
            try:
                data = json.loads(match.group())
            except json.JSONDecodeError as err:
                raise ValueError(f"AI response could not be parsed as JSON.\nParser error: {err}\nRaw (first 500 chars): {clean[:500]}")
        else:
            raise ValueError(f"AI response did not contain a valid JSON object.\nRaw (first 500 chars): {clean[:500]}")

    return _apply_defaults(data)


def _apply_defaults(data: dict) -> dict:
    scalar_defaults = {
        "candidate_name": "Not Specified",
        "email": None, "phone": None, "location": None,
        "linkedin": None, "github": None, "portfolio": None,
        "professional_summary": "No summary provided.",
        "overall_score": 50, "experience_years": 0,
        "education_level": "Not Specified", "education_field": "Not Specified",
        "institution_name": "Not Specified", "projects_count": 0,
        "hire_recommendation": "MAYBE", "hire_confidence": 50,
        "hire_reasoning": "Analysis incomplete.",
    }
    score_defaults = {
        "technical_skills": 50, "cybersecurity_relevance": 50,
        "experience": 50, "education": 50, "soft_skills": 50, "ats_compatibility": 50,
    }
    list_fields = [
        "technical_skills_found", "cybersecurity_skills_found", "soft_skills_found",
        "certifications_found", "tools_found", "programming_languages",
        "required_skills_missing", "recommended_certifications",
        "red_flags", "green_flags", "interview_questions",
    ]

    for key, val in scalar_defaults.items():
        data.setdefault(key, val)

    if not isinstance(data.get("scores"), dict):
        data["scores"] = {}
    for k, v in score_defaults.items():
        raw = data["scores"].get(k, v)
        try:
            data["scores"][k] = max(0, min(100, int(raw)))
        except (TypeError, ValueError):
            data["scores"][k] = v

    try:
        data["overall_score"] = max(0, min(100, int(data["overall_score"])))
    except (TypeError, ValueError):
        vals = list(data["scores"].values())
        data["overall_score"] = int(sum(vals) / len(vals)) if vals else 50

    for field in list_fields:
        if not isinstance(data.get(field), list):
            data[field] = []

    for field in ("strengths", "weaknesses"):
        if not isinstance(data.get(field), list):
            data[field] = []
        data[field] = [
            item for item in data[field]
            if isinstance(item, dict) and "title" in item and "description" in item
        ]

    if not isinstance(data.get("improvements"), list):
        data["improvements"] = []
    data["improvements"] = [
        item for item in data["improvements"]
        if isinstance(item, dict) and "suggestion" in item
    ]
    for item in data["improvements"]:
        item.setdefault("priority", "MEDIUM")
        item.setdefault("area", "General")
        item.setdefault("current_state", "")

    rec = str(data.get("hire_recommendation", "MAYBE")).upper()
    data["hire_recommendation"] = rec if rec in ("HIRE", "MAYBE", "REJECT") else "MAYBE"

    return data


# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 4 — PLOTLY CHARTS
# ═══════════════════════════════════════════════════════════════════════════════

C = {
    "primary": "#00d4ff", "accent": "#00ffa3", "warning": "#ffb800",
    "danger": "#ff3b3b", "benchmark": "#ff6b6b",
    "bg": "rgba(0,0,0,0)", "card_bg": "rgba(15,22,41,0.8)",
    "grid": "#1e2d4d", "text": "#e8eaf6", "muted": "#6b7a99",
}

SCORE_META = {
    "technical_skills":       {"label": "Technical Skills", "benchmark": 68},
    "cybersecurity_relevance":{"label": "Cybersecurity",    "benchmark": 65},
    "experience":             {"label": "Experience",        "benchmark": 70},
    "education":              {"label": "Education",         "benchmark": 65},
    "soft_skills":            {"label": "Soft Skills",       "benchmark": 60},
    "ats_compatibility":      {"label": "ATS Score",         "benchmark": 72},
}


def _base_layout(**overrides) -> dict:
    layout = dict(
        paper_bgcolor=C["bg"],
        plot_bgcolor=C["card_bg"],
        font=dict(color=C["text"], family="'JetBrains Mono', 'Courier New', monospace"),
        margin=dict(l=20, r=20, t=30, b=20),
    )
    layout.update(overrides)
    return layout


def create_radar_chart(scores: dict) -> go.Figure:
    labels = [SCORE_META[k]["label"] for k in SCORE_META]
    cand   = [scores.get(k, 0) for k in SCORE_META]
    bench  = [SCORE_META[k]["benchmark"] for k in SCORE_META]
    lc, cc, bc = labels + [labels[0]], cand + [cand[0]], bench + [bench[0]]

    fig = go.Figure()
    fig.add_trace(go.Scatterpolar(r=bc, theta=lc, fill="toself",
        fillcolor="rgba(255,107,107,0.07)", line=dict(color=C["benchmark"], width=1.5, dash="dot"),
        name="Industry Avg", hovertemplate="%{theta}: %{r}/100<extra>Benchmark</extra>"))
    fig.add_trace(go.Scatterpolar(r=cc, theta=lc, fill="toself",
        fillcolor="rgba(0,212,255,0.15)", line=dict(color=C["primary"], width=2.5),
        name="Candidate", hovertemplate="%{theta}: %{r}/100<extra>Candidate</extra>"))
    fig.update_layout(
        **_base_layout(height=420, margin=dict(l=60, r=60, t=20, b=20)),
        polar=dict(
            bgcolor="rgba(8,12,24,0.6)",
            radialaxis=dict(visible=True, range=[0, 100],
                tickfont=dict(size=9, color=C["muted"]), gridcolor=C["grid"], linecolor=C["grid"]),
            angularaxis=dict(tickfont=dict(size=11, color=C["text"]), gridcolor=C["grid"], linecolor=C["grid"]),
        ),
        legend=dict(font=dict(size=11, color=C["text"]), bgcolor="rgba(15,22,41,0.9)",
            bordercolor=C["grid"], borderwidth=1),
        showlegend=True,
    )
    return fig


def create_overall_gauge(score: int) -> go.Figure:
    if score >= 70:
        bar_color, zone_label = C["accent"], "HIRE ZONE"
    elif score >= 50:
        bar_color, zone_label = C["warning"], "REVIEW ZONE"
    else:
        bar_color, zone_label = C["danger"], "REJECT ZONE"

    fig = go.Figure(go.Indicator(
        mode="gauge+number", value=score,
        number={"font": {"size": 52, "color": bar_color}, "suffix": ""},
        gauge={
            "axis": {"range": [0, 100], "tickfont": {"color": C["muted"], "size": 10},
                "tickwidth": 1, "tickcolor": C["grid"]},
            "bar": {"color": bar_color, "thickness": 0.72},
            "bgcolor": "rgba(15,22,41,0.8)", "borderwidth": 0,
            "steps": [
                {"range": [0,  50], "color": "rgba(255,59,59,0.12)"},
                {"range": [50, 70], "color": "rgba(255,184,0,0.10)"},
                {"range": [70,100], "color": "rgba(0,255,163,0.10)"},
            ],
            "threshold": {"line": {"color": C["text"], "width": 2}, "thickness": 0.82, "value": 70},
        },
        title={"text": f"Overall Score<br><span style='font-size:0.7em;color:{bar_color}'>{zone_label}</span>",
               "font": {"size": 14, "color": C["text"]}},
        domain={"x": [0, 1], "y": [0, 1]},
    ))
    fig.update_layout(**_base_layout(height=280, margin=dict(l=30, r=30, t=40, b=10)))
    return fig


def create_skills_bar_chart(analysis: dict) -> go.Figure:
    categories = {
        "Technical Skills":  len(analysis.get("technical_skills_found", [])),
        "Cyber / Security":  len(analysis.get("cybersecurity_skills_found", [])),
        "Tools & Platforms": len(analysis.get("tools_found", [])),
        "Programming Langs": len(analysis.get("programming_languages", [])),
        "Soft Skills":       len(analysis.get("soft_skills_found", [])),
        "Certifications":    len(analysis.get("certifications_found", [])),
    }
    sorted_cats = dict(sorted(categories.items(), key=lambda x: x[1], reverse=True))
    bar_colors = [C["primary"], C["accent"], "#a78bfa", "#f97316", "#ec4899", C["warning"]]

    fig = go.Figure(go.Bar(
        x=list(sorted_cats.values()), y=list(sorted_cats.keys()), orientation="h",
        marker=dict(color=bar_colors[:len(sorted_cats)], line=dict(color="rgba(0,0,0,0)", width=0)),
        text=[str(v) for v in sorted_cats.values()], textposition="outside",
        textfont=dict(color=C["text"], size=13),
        hovertemplate="%{y}: %{x} items<extra></extra>",
    ))
    fig.update_layout(
        **_base_layout(height=310, margin=dict(l=10, r=50, t=10, b=10)),
        xaxis=dict(gridcolor=C["grid"], tickfont=dict(color=C["muted"]), showgrid=True, zeroline=False),
        yaxis=dict(tickfont=dict(color=C["text"], size=12), gridcolor="rgba(0,0,0,0)"),
        showlegend=False,
    )
    return fig


def create_score_comparison_bar(scores: dict) -> go.Figure:
    labels     = [SCORE_META[k]["label"] for k in SCORE_META]
    cand_vals  = [scores.get(k, 0) for k in SCORE_META]
    bench_vals = [SCORE_META[k]["benchmark"] for k in SCORE_META]

    fig = go.Figure()
    fig.add_trace(go.Bar(name="Candidate", x=labels, y=cand_vals, marker_color=C["primary"],
        text=cand_vals, textposition="outside", textfont=dict(color=C["primary"], size=11),
        hovertemplate="%{x}: %{y}/100<extra>Candidate</extra>"))
    fig.add_trace(go.Bar(name="Industry Avg", x=labels, y=bench_vals,
        marker_color="rgba(255,107,107,0.45)", text=bench_vals, textposition="outside",
        textfont=dict(color=C["benchmark"], size=11),
        hovertemplate="%{x}: %{y}/100<extra>Benchmark</extra>"))
    fig.update_layout(
        **_base_layout(height=340, margin=dict(l=10, r=10, t=10, b=60)),
        barmode="group",
        xaxis=dict(tickfont=dict(color=C["text"], size=10), gridcolor="rgba(0,0,0,0)"),
        yaxis=dict(range=[0, 115], gridcolor=C["grid"], tickfont=dict(color=C["muted"])),
        legend=dict(font=dict(color=C["text"]), bgcolor="rgba(15,22,41,0.9)",
            bordercolor=C["grid"], borderwidth=1),
    )
    return fig


def create_improvement_donut(improvements: list):
    counts = {"HIGH": 0, "MEDIUM": 0, "LOW": 0}
    for item in improvements:
        p = str(item.get("priority", "MEDIUM")).upper()
        if p in counts:
            counts[p] += 1
    total = sum(counts.values())
    if total == 0:
        return None

    fig = go.Figure(go.Pie(
        labels=list(counts.keys()), values=list(counts.values()), hole=0.60,
        marker=dict(colors=[C["danger"], C["warning"], C["accent"]], line=dict(color="#080c18", width=2)),
        textinfo="label+percent", textfont=dict(size=12, color="white"),
        hovertemplate="%{label}: %{value} items<extra></extra>",
    ))
    fig.update_layout(
        **_base_layout(height=240, margin=dict(l=10, r=10, t=10, b=10)),
        showlegend=False,
        annotations=[{"text": f"<b>{total}</b><br>Items", "x": 0.5, "y": 0.5,
            "font": {"size": 16, "color": C["text"]}, "showarrow": False}],
    )
    return fig


# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 5 — STREAMLIT UI
# ═══════════════════════════════════════════════════════════════════════════════

st.set_page_config(
    page_title="AI Resume Analyser",
    page_icon="🧠",
    layout="wide",
    initial_sidebar_state="expanded",
)


def inject_css():
    st.markdown(
        """
        <style>
        @import url('https://fonts.googleapis.com/css2?family=JetBrains+Mono:wght@400;600;700&family=Sora:wght@400;600;700;800&display=swap');

        html, body, [class*="css"] { font-family: 'Sora', sans-serif; }

        ::-webkit-scrollbar { width: 6px; }
        ::-webkit-scrollbar-track { background: #080c18; }
        ::-webkit-scrollbar-thumb { background: #1e2d4d; border-radius: 3px; }

        #MainMenu, footer, header { visibility: hidden; }

        .main .block-container { padding-top: 1.5rem; padding-bottom: 2rem; max-width: 1300px; }

        [data-testid="stSidebar"] { background: #080c18; border-right: 1px solid #1e2d4d; }

        .ra-card {
            background: linear-gradient(135deg, #0f1629 0%, #0d1b36 100%);
            border: 1px solid #1e2d4d; border-radius: 12px;
            padding: 1.2rem 1.4rem; margin-bottom: 0.9rem;
        }
        .ra-card-accent { border-left: 3px solid #00d4ff; }

        .cand-card {
            background: linear-gradient(135deg, #0f1629 0%, #0d1b36 100%);
            border: 1px solid #1e2d4d; border-radius: 14px;
            padding: 1.5rem 2rem; margin-bottom: 1.2rem;
            position: relative; overflow: hidden;
        }
        .cand-card::before {
            content: ''; position: absolute; top: 0; left: 0;
            width: 4px; height: 100%;
            background: linear-gradient(180deg, #00d4ff, #00ffa3); border-radius: 2px;
        }
        .cand-name { font-family: 'Sora', sans-serif; font-size: 1.7rem; font-weight: 800; color: #e8eaf6; margin: 0 0 0.3rem 0; }
        .cand-role { font-size: 0.85rem; color: #00d4ff; font-family: 'JetBrains Mono', monospace; letter-spacing: 1px; text-transform: uppercase; }
        .cand-meta { display: flex; flex-wrap: wrap; gap: 0.8rem; margin-top: 0.9rem; }
        .cand-chip {
            display: inline-flex; align-items: center; gap: 0.35rem;
            background: rgba(0,212,255,0.08); border: 1px solid rgba(0,212,255,0.2);
            border-radius: 20px; padding: 0.25rem 0.7rem; font-size: 0.8rem; color: #b0c4de;
        }
        .cand-chip a { color: #00d4ff; text-decoration: none; }

        .score-grid { display: grid; grid-template-columns: repeat(3, 1fr); gap: 0.7rem; margin-bottom: 1.2rem; }
        .score-tile { background: #0f1629; border: 1px solid #1e2d4d; border-radius: 10px; padding: 0.85rem 1rem; text-align: center; }
        .score-tile-label { font-size: 0.7rem; text-transform: uppercase; letter-spacing: 1px; color: #6b7a99; font-family: 'JetBrains Mono', monospace; margin-bottom: 0.3rem; }
        .score-tile-value { font-size: 1.6rem; font-weight: 800; font-family: 'JetBrains Mono', monospace; }
        .score-tile-bar { height: 3px; border-radius: 2px; margin-top: 0.4rem; background: #1e2d4d; }
        .score-tile-fill { height: 100%; border-radius: 2px; }

        .tag-row { display: flex; flex-wrap: wrap; gap: 0.4rem; margin: 0.5rem 0; }
        .tag { display: inline-block; padding: 3px 10px; border-radius: 20px; font-size: 0.78rem; font-family: 'JetBrains Mono', monospace; font-weight: 600; }
        .tag-tech    { background:#0a2040; color:#60b4ff; border:1px solid #1a4070; }
        .tag-cyber   { background:#0a2818; color:#4ded9e; border:1px solid #1a5c38; }
        .tag-tool    { background:#1e1040; color:#c084fc; border:1px solid #3d2080; }
        .tag-cert    { background:#2a1800; color:#fbbf24; border:1px solid #5a3800; }
        .tag-lang    { background:#1a0a2e; color:#f472b6; border:1px solid #3c1a5c; }
        .tag-soft    { background:#0a1e28; color:#67e8f9; border:1px solid #1a4858; }
        .tag-missing { background:#2a0a0a; color:#f87171; border:1px solid #5a1a1a; }

        .sw-card { border-radius: 10px; padding: 1rem 1.2rem; margin-bottom: 0.7rem; }
        .sw-strength { background: rgba(0,255,163,0.05); border-left: 3px solid #00ffa3; }
        .sw-weakness { background: rgba(255,59,59,0.05);  border-left: 3px solid #ff3b3b; }
        .sw-title { font-weight: 700; font-size: 0.92rem; margin-bottom: 0.3rem; }
        .sw-desc  { font-size: 0.85rem; color: #9ba8c0; line-height: 1.55; }

        .flag-row { display: flex; flex-wrap: wrap; gap: 0.5rem; margin: 0.5rem 0; }
        .flag { padding: 4px 12px; border-radius: 20px; font-size: 0.8rem; font-weight: 600; }
        .flag-green { background:rgba(0,255,163,0.1); color:#00ffa3; border:1px solid rgba(0,255,163,0.2); }
        .flag-red   { background:rgba(255,59,59,0.1);  color:#ff3b3b; border:1px solid rgba(255,59,59,0.2);  }

        .imp-card { border-radius: 10px; padding: 1rem 1.2rem; margin-bottom: 0.7rem; }
        .imp-HIGH   { background:rgba(255,59,59,0.05);  border-left:3px solid #ff3b3b; }
        .imp-MEDIUM { background:rgba(255,184,0,0.05);  border-left:3px solid #ffb800; }
        .imp-LOW    { background:rgba(0,255,163,0.05);  border-left:3px solid #00ffa3; }
        .imp-badge { display:inline-block; padding:2px 8px; border-radius:4px; font-size:0.68rem; font-family:'JetBrains Mono',monospace; font-weight:700; letter-spacing:1px; margin-bottom:0.5rem; }
        .badge-HIGH   { background:#3d0f0f; color:#ff3b3b; }
        .badge-MEDIUM { background:#3d2e00; color:#ffb800; }
        .badge-LOW    { background:#0a2818; color:#00ffa3; }
        .imp-area       { font-weight:700; font-size:0.9rem; margin-bottom:0.3rem; }
        .imp-current    { font-size:0.8rem; color:#6b7a99; margin-bottom:0.4rem; }
        .imp-suggestion { font-size:0.85rem; color:#9ba8c0; line-height:1.55; }

        .hire-card { border-radius: 16px; padding: 2.5rem 2rem; text-align: center; margin-bottom: 1.5rem; }
        .hire-HIRE   { background: linear-gradient(135deg,#001f12,#003d22); border:2px solid #00ffa3; box-shadow:0 0 40px rgba(0,255,163,0.15); }
        .hire-MAYBE  { background: linear-gradient(135deg,#1f1600,#3d2d00); border:2px solid #ffb800; box-shadow:0 0 40px rgba(255,184,0,0.15); }
        .hire-REJECT { background: linear-gradient(135deg,#1f0000,#3d0000); border:2px solid #ff3b3b; box-shadow:0 0 40px rgba(255,59,59,0.15); }
        .hire-icon      { font-size: 3.5rem; margin-bottom: 0.5rem; }
        .hire-verdict   { font-family: 'JetBrains Mono', monospace; font-size: 2.2rem; font-weight: 800; letter-spacing: 4px; }
        .hire-confidence{ font-size: 0.9rem; color: #9ba8c0; margin-top: 0.5rem; }

        .iq-card { background:rgba(0,212,255,0.04); border:1px solid rgba(0,212,255,0.12); border-radius:10px; padding:1rem 1.2rem; margin-bottom:0.6rem; }
        .iq-num  { font-family:'JetBrains Mono',monospace; font-size:0.75rem; color:#00d4ff; font-weight:700; margin-bottom:0.3rem; }
        .iq-text { font-size:0.88rem; color:#c0cce8; line-height:1.55; }

        .section-heading { font-family:'JetBrains Mono',monospace; font-size:0.72rem; text-transform:uppercase; letter-spacing:2px; color:#00d4ff; border-bottom:1px solid #1e2d4d; padding-bottom:0.5rem; margin:1.2rem 0 0.8rem 0; }

        [data-testid="stFileUploader"] { border:2px dashed #1e2d4d !important; border-radius:12px !important; background:#0a0f20 !important; }

        .stTabs [data-baseweb="tab-list"] { background:#0f1629; border-radius:10px; padding:4px; gap:6px; border:1px solid #1e2d4d; }
        .stTabs [data-baseweb="tab"]      { border-radius:8px; color:#6b7a99; font-size:0.85rem; font-family:'JetBrains Mono',monospace; }
        .stTabs [aria-selected="true"]    { background:#00d4ff !important; color:#080c18 !important; font-weight:700; }

        .sidebar-logo { text-align:center; padding:1rem 0 1.5rem 0; }
        .sidebar-logo-text { font-family:'JetBrains Mono',monospace; font-size:1.1rem; font-weight:700; color:#00d4ff; letter-spacing:2px; }
        .sidebar-tagline   { font-size:0.72rem; color:#6b7a99; margin-top:0.2rem; }

        .stat-bar-row   { display:flex; align-items:center; gap:0.8rem; margin-bottom:0.55rem; }
        .stat-bar-label { font-size:0.78rem; color:#9ba8c0; width:140px; flex-shrink:0; font-family:'JetBrains Mono',monospace; }
        .stat-bar-track { flex:1; height:6px; background:#1e2d4d; border-radius:3px; overflow:hidden; }
        .stat-bar-fill  { height:100%; border-radius:3px; }
        .stat-bar-val   { font-size:0.78rem; font-weight:700; width:32px; text-align:right; font-family:'JetBrains Mono',monospace; }
        </style>
        """,
        unsafe_allow_html=True,
    )


# ─── HTML helpers ─────────────────────────────────────────────────────────────

def score_color(val: int) -> str:
    if val >= 70:
        return "#00ffa3"
    if val >= 50:
        return "#ffb800"
    return "#ff3b3b"


def render_candidate_card(analysis: dict, job_role: str) -> None:
    name      = analysis.get("candidate_name", "Unknown Candidate")
    email     = analysis.get("email")
    phone     = analysis.get("phone")
    location  = analysis.get("location")
    linkedin  = analysis.get("linkedin")
    github    = analysis.get("github")
    portfolio = analysis.get("portfolio")
    edu       = analysis.get("education_level", "")
    edu_field = analysis.get("education_field", "")
    exp_yrs   = analysis.get("experience_years", 0)
    proj      = analysis.get("projects_count", 0)

    chips = ""
    if email:     chips += f'<span class="cand-chip">✉ {email}</span>'
    if phone:     chips += f'<span class="cand-chip">📞 {phone}</span>'
    if location:  chips += f'<span class="cand-chip">📍 {location}</span>'
    if linkedin:  chips += f'<span class="cand-chip"><a href="{linkedin}" target="_blank">🔗 LinkedIn</a></span>'
    if github:    chips += f'<span class="cand-chip"><a href="{github}" target="_blank">💻 GitHub</a></span>'
    if portfolio: chips += f'<span class="cand-chip"><a href="{portfolio}" target="_blank">🌐 Portfolio</a></span>'
    if edu and edu != "Not Specified": chips += f'<span class="cand-chip">🎓 {edu} — {edu_field}</span>'
    if exp_yrs:   chips += f'<span class="cand-chip">⏱ {exp_yrs} yr exp</span>'
    if proj:      chips += f'<span class="cand-chip">🗂 {proj} projects</span>'

    st.markdown(
        f'<div class="cand-card">'
        f'<div class="cand-name">{name}</div>'
        f'<div class="cand-role">→ Applied for: {job_role}</div>'
        f'<div class="cand-meta">{chips}</div>'
        f'</div>',
        unsafe_allow_html=True,
    )


def render_tag_row(items: list, css_class: str) -> None:
    if not items:
        st.markdown('<span style="color:#6b7a99;font-size:0.82rem;">None found</span>', unsafe_allow_html=True)
        return
    html = '<div class="tag-row">' + "".join(f'<span class="tag {css_class}">{i}</span>' for i in items) + "</div>"
    st.markdown(html, unsafe_allow_html=True)


def render_sw_cards(items: list, card_class: str) -> None:
    for item in items:
        st.markdown(
            f'<div class="sw-card {card_class}">'
            f'<div class="sw-title">{item.get("title","")}</div>'
            f'<div class="sw-desc">{item.get("description","")}</div>'
            f'</div>',
            unsafe_allow_html=True,
        )


def render_improvement_cards(improvements: list) -> None:
    order = {"HIGH": 0, "MEDIUM": 1, "LOW": 2}
    for item in sorted(improvements, key=lambda x: order.get(str(x.get("priority","MEDIUM")).upper(), 1)):
        p       = str(item.get("priority", "MEDIUM")).upper()
        area    = item.get("area", "General")
        current = item.get("current_state", "")
        tip     = item.get("suggestion", "")
        st.markdown(
            f'<div class="imp-card imp-{p}">'
            f'<span class="imp-badge badge-{p}">{p} PRIORITY</span>'
            f'<div class="imp-area">{area}</div>'
            + (f'<div class="imp-current">Currently: {current}</div>' if current else "") +
            f'<div class="imp-suggestion">→ {tip}</div>'
            f'</div>',
            unsafe_allow_html=True,
        )


def render_flags(green_flags: list, red_flags: list) -> None:
    c1, c2 = st.columns(2)
    with c1:
        st.markdown('<div class="section-heading">✅ Green Flags</div>', unsafe_allow_html=True)
        if green_flags:
            st.markdown('<div class="flag-row">' + "".join(f'<span class="flag flag-green">✓ {f}</span>' for f in green_flags) + "</div>", unsafe_allow_html=True)
        else:
            st.markdown('<span style="color:#6b7a99;font-size:0.82rem;">None identified</span>', unsafe_allow_html=True)
    with c2:
        st.markdown('<div class="section-heading">🚩 Red Flags</div>', unsafe_allow_html=True)
        if red_flags:
            st.markdown('<div class="flag-row">' + "".join(f'<span class="flag flag-red">⚠ {f}</span>' for f in red_flags) + "</div>", unsafe_allow_html=True)
        else:
            st.markdown('<span style="color:#6b7a99;font-size:0.82rem;">None identified</span>', unsafe_allow_html=True)


def render_stat_bar(label: str, val: int) -> None:
    color = score_color(val)
    st.markdown(
        f'<div class="stat-bar-row">'
        f'<div class="stat-bar-label">{label}</div>'
        f'<div class="stat-bar-track"><div class="stat-bar-fill" style="width:{val}%;background:{color};"></div></div>'
        f'<div class="stat-bar-val" style="color:{color};">{val}</div>'
        f'</div>',
        unsafe_allow_html=True,
    )


# ─── Report generator ─────────────────────────────────────────────────────────

def generate_text_report(analysis: dict, job_role: str) -> str:
    scores = analysis.get("scores", {})
    score_labels = {
        "technical_skills": "Technical Skills", "cybersecurity_relevance": "Cybersecurity    ",
        "experience": "Experience       ", "education": "Education        ",
        "soft_skills": "Soft Skills      ", "ats_compatibility": "ATS Compatibility",
    }
    lines = [
        "=" * 70,
        "          AI RESUME ANALYSIS REPORT",
        f"          Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
        f"          Role: {job_role}",
        "=" * 70, "",
        "CANDIDATE", "-" * 40,
        f"Name     : {analysis.get('candidate_name','N/A')}",
        f"Email    : {analysis.get('email') or 'N/A'}",
        f"Phone    : {analysis.get('phone') or 'N/A'}",
        f"Location : {analysis.get('location') or 'N/A'}",
        f"LinkedIn : {analysis.get('linkedin') or 'N/A'}",
        f"GitHub   : {analysis.get('github') or 'N/A'}",
        f"Exp Yrs  : {analysis.get('experience_years',0)}",
        f"Education: {analysis.get('education_level','N/A')} — {analysis.get('education_field','N/A')}",
        f"Projects : {analysis.get('projects_count',0)} mentioned",
        "", "SUMMARY", "-" * 40,
        analysis.get("professional_summary", "N/A"),
        "", "SCORES", "-" * 40,
    ]
    for key, label in score_labels.items():
        val = scores.get(key, 0)
        bar = "█" * (val // 5) + "░" * (20 - val // 5)
        lines.append(f"{label}: [{bar}] {val}/100")

    lines += [f"\nOVERALL SCORE: {analysis.get('overall_score',0)}/100", "", "SKILLS FOUND", "-" * 40,
        f"Technical  : {', '.join(analysis.get('technical_skills_found',[])) or 'None'}",
        f"Cybersec   : {', '.join(analysis.get('cybersecurity_skills_found',[])) or 'None'}",
        f"Tools      : {', '.join(analysis.get('tools_found',[])) or 'None'}",
        f"Languages  : {', '.join(analysis.get('programming_languages',[])) or 'None'}",
        f"Certs      : {', '.join(analysis.get('certifications_found',[])) or 'None'}",
        f"Missing    : {', '.join(analysis.get('required_skills_missing',[])) or 'None'}",
        "", "STRENGTHS", "-" * 40]
    for s in analysis.get("strengths", []):
        lines.append(f"[+] {s.get('title','')}: {s.get('description','')}")
    lines += ["", "WEAKNESSES", "-" * 40]
    for w in analysis.get("weaknesses", []):
        lines.append(f"[-] {w.get('title','')}: {w.get('description','')}")
    lines += ["", "IMPROVEMENTS", "-" * 40]
    for imp in analysis.get("improvements", []):
        lines.append(f"[{imp.get('priority','MED')}] {imp.get('area','')}: {imp.get('suggestion','')}")
    lines += ["", "GREEN FLAGS", "-" * 40]
    for f in analysis.get("green_flags", []):
        lines.append(f"  ✓ {f}")
    lines += ["", "RED FLAGS", "-" * 40]
    for f in analysis.get("red_flags", []):
        lines.append(f"  ⚠ {f}")
    lines += ["", "INTERVIEW QUESTIONS", "-" * 40]
    for i, q in enumerate(analysis.get("interview_questions", []), 1):
        lines.append(f"  Q{i}: {q}")

    rec = analysis.get("hire_recommendation", "MAYBE")
    lines += [
        "", "=" * 70,
        f"  HIRE DECISION: {rec}  (Confidence: {analysis.get('hire_confidence',50)}%)",
        "=" * 70, "", "Reasoning:", analysis.get("hire_reasoning", ""), "",
        "─" * 70, "Report generated by AI Resume Analyser | github.com/Tktirth", "─" * 70,
    ]
    return "\n".join(lines)


# ─── Tab renderers ────────────────────────────────────────────────────────────

def render_tab_overview(analysis: dict, job_role: str) -> None:
    render_candidate_card(analysis, job_role)
    col_gauge, col_radar = st.columns([1, 1.4], gap="medium")

    with col_gauge:
        st.plotly_chart(create_overall_gauge(analysis.get("overall_score", 0)),
            use_container_width=True, config={"displayModeBar": False})
        st.markdown('<div class="section-heading">Dimension Scores</div>', unsafe_allow_html=True)
        for key, label in {
            "technical_skills": "Technical Skills", "cybersecurity_relevance": "Cybersecurity",
            "experience": "Experience", "education": "Education",
            "soft_skills": "Soft Skills", "ats_compatibility": "ATS Score",
        }.items():
            render_stat_bar(label, analysis["scores"].get(key, 0))

    with col_radar:
        st.markdown('<div class="section-heading">Radar — Candidate vs Industry Avg</div>', unsafe_allow_html=True)
        st.plotly_chart(create_radar_chart(analysis.get("scores", {})),
            use_container_width=True, config={"displayModeBar": False})

    st.markdown('<div class="section-heading">Score Breakdown vs Benchmark</div>', unsafe_allow_html=True)
    st.plotly_chart(create_score_comparison_bar(analysis.get("scores", {})),
        use_container_width=True, config={"displayModeBar": False})
    st.markdown('<div class="section-heading">Professional Summary</div>', unsafe_allow_html=True)
    st.markdown(
        f'<div class="ra-card ra-card-accent" style="font-size:0.9rem;line-height:1.7;color:#c0cce8;">'
        f'{analysis.get("professional_summary","")}</div>',
        unsafe_allow_html=True,
    )


def render_tab_skills(analysis: dict) -> None:
    col_chart, col_right = st.columns([1, 1.3], gap="medium")
    with col_chart:
        st.markdown('<div class="section-heading">Skill Categories Found</div>', unsafe_allow_html=True)
        st.plotly_chart(create_skills_bar_chart(analysis),
            use_container_width=True, config={"displayModeBar": False})
    with col_right:
        st.markdown('<div class="section-heading">Technical Skills</div>', unsafe_allow_html=True)
        render_tag_row(analysis.get("technical_skills_found", []), "tag-tech")
        st.markdown('<div class="section-heading">Cybersecurity Skills</div>', unsafe_allow_html=True)
        render_tag_row(analysis.get("cybersecurity_skills_found", []), "tag-cyber")
        st.markdown('<div class="section-heading">Programming Languages</div>', unsafe_allow_html=True)
        render_tag_row(analysis.get("programming_languages", []), "tag-lang")

    c1, c2 = st.columns(2, gap="medium")
    with c1:
        st.markdown('<div class="section-heading">Tools & Platforms</div>', unsafe_allow_html=True)
        render_tag_row(analysis.get("tools_found", []), "tag-tool")
    with c2:
        st.markdown('<div class="section-heading">Certifications</div>', unsafe_allow_html=True)
        render_tag_row(analysis.get("certifications_found", []), "tag-cert")

    st.markdown('<div class="section-heading">⚠ Required Skills Missing for This Role</div>', unsafe_allow_html=True)
    render_tag_row(analysis.get("required_skills_missing", []), "tag-missing")
    if analysis.get("recommended_certifications"):
        st.markdown('<div class="section-heading">📌 Recommended Certifications</div>', unsafe_allow_html=True)
        render_tag_row(analysis.get("recommended_certifications", []), "tag-cert")


def render_tab_strengths(analysis: dict) -> None:
    c1, c2 = st.columns(2, gap="medium")
    with c1:
        st.markdown('<div class="section-heading">💪 Strengths</div>', unsafe_allow_html=True)
        strengths = analysis.get("strengths", [])
        render_sw_cards(strengths, "sw-strength") if strengths else st.info("No strengths identified.")
    with c2:
        st.markdown('<div class="section-heading">⚡ Weaknesses</div>', unsafe_allow_html=True)
        weaknesses = analysis.get("weaknesses", [])
        render_sw_cards(weaknesses, "sw-weakness") if weaknesses else st.info("No weaknesses identified.")
    render_flags(analysis.get("green_flags", []), analysis.get("red_flags", []))


def render_tab_improvements(analysis: dict) -> None:
    improvements = analysis.get("improvements", [])
    c1, c2 = st.columns([1, 2.2], gap="medium")
    with c1:
        st.markdown('<div class="section-heading">Priority Distribution</div>', unsafe_allow_html=True)
        donut = create_improvement_donut(improvements)
        if donut:
            st.plotly_chart(donut, use_container_width=True, config={"displayModeBar": False})
        else:
            st.info("No improvements recorded.")
    with c2:
        st.markdown('<div class="section-heading">Action Plan</div>', unsafe_allow_html=True)
        if improvements:
            render_improvement_cards(improvements)
        else:
            st.success("No specific improvements required. Strong candidate.")

    qs = analysis.get("interview_questions", [])
    if qs:
        st.markdown('<div class="section-heading">🎙 Suggested Interview Questions</div>', unsafe_allow_html=True)
        for i, q in enumerate(qs, 1):
            st.markdown(
                f'<div class="iq-card"><div class="iq-num">QUESTION {i:02d}</div>'
                f'<div class="iq-text">{q}</div></div>',
                unsafe_allow_html=True,
            )


def render_tab_decision(analysis: dict, job_role: str) -> None:
    rec     = analysis.get("hire_recommendation", "MAYBE")
    conf    = analysis.get("hire_confidence", 50)
    reason  = analysis.get("hire_reasoning", "")
    overall = analysis.get("overall_score", 0)

    icons = {"HIRE": "✅", "MAYBE": "⚠️", "REJECT": "❌"}
    vc    = {"HIRE": "#00ffa3", "MAYBE": "#ffb800", "REJECT": "#ff3b3b"}

    st.markdown(
        f'<div class="hire-card hire-{rec}">'
        f'<div class="hire-icon">{icons.get(rec,"⚠️")}</div>'
        f'<div class="hire-verdict" style="color:{vc.get(rec,"#ffb800")};">{rec}</div>'
        f'<div class="hire-confidence">Overall Score: {overall}/100 &nbsp;|&nbsp; Confidence: {conf}%</div>'
        f'</div>',
        unsafe_allow_html=True,
    )
    st.markdown('<div class="section-heading">HR Decision Reasoning</div>', unsafe_allow_html=True)
    st.markdown(
        f'<div class="ra-card ra-card-accent" style="font-size:0.88rem;line-height:1.75;color:#c0cce8;">{reason}</div>',
        unsafe_allow_html=True,
    )
    if rec == "HIRE":
        st.success("**Recommendation:** Proceed to offer. Schedule a technical interview to confirm depth.")
    elif rec == "MAYBE":
        st.warning("**Recommendation:** Schedule a screening call. Address flagged gaps before final decision.")
    else:
        st.error("**Recommendation:** Do not proceed. Candidate does not meet minimum requirements for this role.")

    st.markdown("---")
    name_slug = analysis.get("candidate_name", "candidate").replace(" ", "_")
    st.download_button("⬇  Download Full Report (.txt)", generate_text_report(analysis, job_role),
        file_name=f"resume_analysis_{name_slug}.txt", mime="text/plain", use_container_width=True)
    st.download_button("⬇  Download Raw JSON", json.dumps(analysis, indent=2),
        file_name=f"resume_analysis_{name_slug}.json", mime="application/json", use_container_width=True)


# ─── Sidebar ──────────────────────────────────────────────────────────────────

def render_sidebar() -> tuple:
    with st.sidebar:
        st.markdown(
            '<div class="sidebar-logo">'
            '<div class="sidebar-logo-text">[ AI·RESUME·HR ]</div>'
            '<div class="sidebar-tagline">Powered by LLaMA 3.3 · 70B</div>'
            '</div>',
            unsafe_allow_html=True,
        )
        st.divider()

        api_key = st.text_input("🔑 Groq API Key",
            value=os.getenv("GROQ_API_KEY", ""), type="password",
            placeholder="gsk_...", help="Free key at console.groq.com")

        st.markdown("---")

        job_role  = st.selectbox("🎯 Hiring For", options=get_role_list(), index=0,
            help="Select the role you are currently hiring for.")
        role_info = get_role_info(job_role)
        if role_info.get("description"):
            st.caption(f"{role_info['icon']}  {role_info['description']}")

        st.markdown("---")

        with st.expander("📄 Paste Job Description (optional)", expanded=False):
            job_description = st.text_area("Job Description", height=180,
                placeholder="Paste the full job description here for more accurate analysis...",
                label_visibility="collapsed")

        st.markdown("---")

        if st.button("🗑  Clear Analysis", use_container_width=True):
            for key in ["analysis", "resume_text", "analysed_file"]:
                st.session_state.pop(key, None)
            st.rerun()

        st.markdown("---")
        st.caption("Built by [@Tktirth](https://github.com/Tktirth)")
        st.caption("AI Resume Analyser v2.0")

    return api_key, job_role, job_description


# ─── Main ─────────────────────────────────────────────────────────────────────

def main():
    inject_css()

    for key in ("analysis", "resume_text", "analysed_file"):
        if key not in st.session_state:
            st.session_state[key] = None

    api_key, job_role, job_description = render_sidebar()

    st.markdown(
        '<div style="margin-bottom:1.5rem;">'
        '<div style="font-family:\'JetBrains Mono\',monospace;font-size:0.7rem;color:#00d4ff;letter-spacing:3px;text-transform:uppercase;">'
        'Human Resources · AI Talent Suite</div>'
        '<h1 style="font-family:\'Sora\',sans-serif;font-size:2.4rem;font-weight:800;color:#e8eaf6;margin:0.2rem 0;">'
        'Resume Analyser</h1>'
        '<p style="color:#6b7a99;font-size:0.88rem;margin:0;">'
        'Upload a candidate\'s resume. Get AI-powered scoring, skill gaps, improvement plans, and a final hire/reject decision — instantly.</p>'
        '</div>',
        unsafe_allow_html=True,
    )

    uploaded_file = st.file_uploader("Drop resume here (PDF, DOCX, or TXT)",
        type=["pdf", "docx", "txt"], help="Max recommended size: 10 MB. Text-based PDFs work best.")

    if uploaded_file is not None:
        file_changed = st.session_state.analysed_file != uploaded_file.name
        col_info, col_btn = st.columns([3, 1], gap="small")
        wc_placeholder = col_info.empty()

        with col_btn:
            run_analysis = st.button("🔍  Analyse Resume", use_container_width=True,
                type="primary", disabled=(not api_key))

        if not api_key:
            st.warning("Enter your Groq API key in the sidebar to proceed.")

        if run_analysis or (st.session_state.analysis is not None and not file_changed):
            if run_analysis:
                with st.spinner("Extracting resume text…"):
                    try:
                        resume_text = extract_text_from_resume(uploaded_file)
                        st.session_state.resume_text = resume_text
                        st.session_state.analysed_file = uploaded_file.name
                    except (ValueError, RuntimeError) as e:
                        st.error(f"**Extraction Failed:** {e}")
                        st.stop()

                wc_placeholder.caption(f"📄  **{uploaded_file.name}** · {get_word_count(st.session_state.resume_text):,} words extracted")

                if len(st.session_state.resume_text.strip()) < 100:
                    st.error("Extracted text is too short. This looks like a blank or image-only file.")
                    st.stop()

                progress = st.progress(0, "Connecting to Groq API…")
                try:
                    progress.progress(20, "Sending resume to LLaMA 3.3 70B…")
                    result = analyze_resume(
                        resume_text=st.session_state.resume_text,
                        job_role=job_role,
                        job_description=job_description,
                        api_key=api_key,
                        role_info=get_role_info(job_role),
                    )
                    progress.progress(90, "Parsing AI response…")
                    st.session_state.analysis = result
                    progress.progress(100, "Done.")
                    progress.empty()
                except (RuntimeError, ValueError) as e:
                    progress.empty()
                    st.error(f"**Analysis Failed:** {e}")
                    st.stop()

            if st.session_state.analysis:
                analysis = st.session_state.analysis
                wc_placeholder.caption(f"📄  **{st.session_state.analysed_file}** · {get_word_count(st.session_state.resume_text or ''):,} words extracted")
                st.markdown("---")

                tab1, tab2, tab3, tab4, tab5 = st.tabs([
                    "📊  Overview", "🛠  Skills",
                    "💪  Strengths & Weaknesses", "📈  Improvements", "✅  Hire Decision",
                ])
                with tab1: render_tab_overview(analysis, job_role)
                with tab2: render_tab_skills(analysis)
                with tab3: render_tab_strengths(analysis)
                with tab4: render_tab_improvements(analysis)
                with tab5: render_tab_decision(analysis, job_role)

                with st.expander("🔎 View Extracted Resume Text", expanded=False):
                    st.text(st.session_state.resume_text[:3000])
                    if len(st.session_state.resume_text) > 3000:
                        st.caption("(Showing first 3,000 characters)")

    else:
        st.markdown("---")
        c1, c2, c3 = st.columns(3, gap="medium")
        for col, (icon, title, desc) in zip([c1, c2, c3], [
            ("🤖", "AI-Powered Analysis", "LLaMA 3.3 70B evaluates skills, experience, education, and ATS compatibility against industry benchmarks."),
            ("🛡️", "Cybersecurity-First", "Deep scoring for security roles: pentesting, SOC, DevSecOps, cloud security, and more."),
            ("✅", "Hire / Reject Decision", "Clear recommendation with confidence score, reasoning, and suggested interview questions."),
        ]):
            with col:
                st.markdown(
                    f'<div class="ra-card" style="text-align:center;padding:1.8rem 1.2rem;">'
                    f'<div style="font-size:2.2rem;margin-bottom:0.7rem;">{icon}</div>'
                    f'<div style="font-weight:700;font-size:0.95rem;margin-bottom:0.5rem;">{title}</div>'
                    f'<div style="color:#6b7a99;font-size:0.82rem;line-height:1.6;">{desc}</div>'
                    f'</div>',
                    unsafe_allow_html=True,
                )


if __name__ == "__main__":
    main()
